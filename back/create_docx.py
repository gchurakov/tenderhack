import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pypandoc
from docx2pdf import convert
from datetime import date
# counter of report and contract
n_report = 1
n_contract = 1


def create_changes_report(changes:dict, path:str='./docx_files/', *args, **kwargs) -> str:
    'create changes report -> output filename'
    global n_report

    # TODO fill gaps from kwargs

    doc = Document()

    h1 = doc.add_heading('ПРОТОКОЛ РАЗНОГЛАСИЙ', level=1)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in h1.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    h1_add = doc.add_paragraph(f'к договору поставки № {kwargs["contract_n"]} от {"contract_date"}')# TODO
    h1_add.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in h1_add.runs:
        run.bold = True

    d = date.today()
    today = '.'.join([str(i) for i in (d.day, d.month, d.year)])

    city_date_paragraph = doc.add_paragraph()
    city_date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    city_date_paragraph.add_run(f'{kwargs["place"]}\t\t\t\t\t\t\t\t\t{today}').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # TODO fill input data f-str
    doc.add_paragraph(f'\t{kwargs["supplier"]}, именуемое в дальнейшем «Заказчик», '
                      f'в лице директора Иванова Ивана Ивановича, действующего на основании Устава, с одной стороны,\n '
                      f'\t{kwargs["contractor"]}, именуемое в дальнейшем «Поставщик», '
                      f'в лице директора Петра Петровича Петрова, действующего на основании Устава, с другой стороны\n'
                      f'\tсоставили настоящий протокол разногласий к договору поставки  № {kwargs["contract_n"]} от kwargs["contract_date"] о нижеследующем:'
                      ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_data = [['Пункт договора', 'Текущая версия', 'Предложенные изменения']] + changes
    changes_table = doc.add_table(rows=len(table_data), cols=3)
    changes_table.autofit = False
    changes_table.style = 'Table Grid'
    changes_table.columns[0].width = Cm(0.5)
    changes_table.columns[0].width = Cm(5)
    changes_table.columns[1].width = Cm(5)

    for r in range(len(table_data)):
        for c in range(len(table_data[0])):
            changes_table.cell(r, c).text = table_data[r][c]

    # Add the signature section
    doc.add_paragraph(
        '\n\t1. Настоящий Протокол разногласий составлен в двух экземплярах, имеющий одинаковую юридическую силу '
        'для обеих сторон.').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add the signature lines
    signature_table = doc.add_table(rows=1, cols=2)
    signature_table.autofit = False
    signature_table.style = 'Table Grid'
    signature_table.columns[0].width = Cm(5)
    signature_table.columns[1].width = Cm(5)


    # TODO КОСТЫЛЬ
    kwargs['contractor_address'] = '*'
    kwargs['contractor_ogrn'] = '*'
    kwargs['contractor_inn'] = '*'
    kwargs['contractor_kpp'] = '*'
    kwargs['contractor_signer'] = '*'
    kwargs['contractor_ogrn'] = '*'
    kwargs['supplier_address'] = '*'
    kwargs['supplier_ogrn'] = '*'
    kwargs['supplier_inn'] = '*'
    kwargs['supplier_kpp'] = '*'
    kwargs['supplier_signer'] = '*'


    signature_table.cell(0, 0).text = f'ЗАКАЗЧИК\n{kwargs["contractor"]}\nЮридический и почтовый адрес: {kwargs["contractor_address"]}\nОГРН {kwargs["contractor_ogrn"]}\n' \
                                   f'ИНН {kwargs["contractor_inn"]}\nКПП {kwargs["contractor_kpp"]}\n{kwargs["contractor_signer"]}\n__________________________/{kwargs["contractor_signer"]}/\nм.п.'

    signature_table.cell(0,1).text = f'ПОДРЯДЧИК\n{kwargs["supplier"]}\nЮридический и почтовый адрес: {kwargs["supplier_address"]}\nОГРН {kwargs["supplier_ogrn"]}\n' \
                                   f'ИНН {kwargs["supplier_inn"]}\nКПП {kwargs["supplier_kpp"]}\n{kwargs["supplier_signer"]}\n__________________________/{kwargs["supplier_signer"]}/\nм.п.'

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Save the document
    filename = f'{path}changes_report_{n_report}.docx'
    n_report += 1
    doc.save(filename)

    return filename


def create_contract(path_to_save: str = './docx_files/', **kwargs: dict) -> str:
    'create docx contract from dict with info -> output filename'

    global n_contract
    doc = Document('./docx_files/contract.docx')

    kwargs['contract_n'] = n_contract

    for p in doc.paragraphs:
        for k, v in kwargs.items():
            p.text = p.text.replace("{"+k+"}", str(v))

    for table in doc.tables:
        for cell in table._cells:
            for k, v in kwargs.items():
                cell.text = cell.text.replace("{"+k+"}", str(v))

    filename = f'{path_to_save}contract_{n_contract}.docx'
    n_contract += 1
    doc.save(filename)

    return filename


def to_pdf(filename_docx:str, filename_pdf:str=None) -> str:
    # 'filename .docx -> convert file to .pdf'
    # filename_pdf = '.' + ''.join(filename_docx.split('.')[:-1]) + '.pdf' if filename_pdf is None else filename_pdf
    # convert_file(filename_docx, 'pdf', outputfile=filename_pdf, extra_args=extra_args)
    # TODO convert to pdf
    convert(filename_docx,filename_pdf)
    return filename_pdf


# TODO clear ROFL before code-review
test_changes = [['3.9999.', 'В соответствии с договором', 'Срок оказания услуг составляет 63 календарных дня (не включает Новогодние, '
                                                     'Рождественские и другие праздничные дни на территории РФ) с момента получения '
                                                     'предоплаты в размере 15% от суммы Договора.'],
           ['п.10.2.', 'Заказчик обязан внести предоплату в размере 25%', 'Заказчик обязан внести предоплату в размере 50%'],
           ['п.1', 'В соответствии с договором', 'Изложено в Приложении №1к настоящему протоколу разногласий Ссылка  на файл']]

test_params = {'details' : "На поставку расходного материала (Сетка-слинг)",
               'place' : "г. Пермь",
               'supplier' : "ООО Pizdets",
               'contractor': "ЗАО Ebis",
               'contract_n': "123"}

report_name = create_changes_report(test_changes, **test_params)
contract_name = create_contract(**test_params)
contract_pdf = to_pdf(contract_name)

