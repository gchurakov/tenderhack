import os
import re
from docx import Document
from docx2pdf import convert
from datetime import datetime
import json
# from .core import *
# from .db import engine, db_session, Base
# from .models.core import Tender, Document, ContractClause
# counter of report and contract
n_report = 1
n_contract = 1
locations = {
    'finance_source' : "1",
    'contractor_address': "Часть 2",
    'supplier': "Преамбула",
    'document_name': "1",
    'contractor_oktmo': "11",
    'ikz': "1",
    'reciver': "11",
    'prepay': "2",
    'contractor_kpp': "11",
    'contractor_okpo': "11",
    'price': "2",
    'place': "Преамбула",
    'supplier_signer_in': "11",
    'contractor_inn': "11",
    'contractor_ogrn': "11",
    'contract_date': "Преамбула",
    'contractor': "Преамбула",
    'supplier_email': "11",
    'supplier_ogrn': "11",
    'period': "3",
    'supplier_kpp': "11",
    'supplier_bik': "11",
    'contactor_bank': "11",
    'contractor_bik': "11",
    'contractor_email': "11",
    'contractor_okato': "11",
    'contractor_bank_k_account': "11",
    'supplier_inn': "11",
    'today': "Преамбула",
    'supplier_signer': "Преамбула",
    'contractor_signer': "Преамбула",
    'price_str': "2",
    'contractor_signer_in': "11",
    'supplier_address': "11",
    'contractorr_phone': "11",
    'contract_n': "Преамбула",
    'supplier_bank': "11",
    'supplier_phone': "11",
    'supplier_bank_k_account': "11"
}


def changes_report_fill(changes: list,  raw_json: dict, dirname: str = '',  old_value=None) -> str:
    'create changes report -> output filename'
    global n_report
    input_filename = './docx_files/changes.docx' if filename is None else filename
    output_filename = f'.{dirname}/docx_files/changes_{n_report}.docx'

    if not os.path.exists(dirname):
        os.mkdir(dirname)

    doc = Document(input_filename)
    # changes = [locations[tag], '', raw_json[value]]
    changes_table = doc.tables[0]
    row = changes_table.add_row()
    row.cells[0].text = changes[0]
    row.cells[1].text = f"{old_value}"
    row.cells[2].text = changes[2]

    # raw_json['contract_n'] = n_contract
    # print(raw_json)
    pattern = r'\{"([^"]+)"="([^"]+)"\}'

    for p in doc.paragraphs:
        p.text = re.sub(pattern, r'\2', p.text)

    for table in doc.tables:
        for cell in table._cells:
            cell.text = re.sub(pattern, r'\2', cell.text)

    n_report += 1
    doc.save(output_filename)
    return filename



def contract_fill(raw_json: dict, dirname : str,  input_filename: str = None,  output_filename: str = None) -> str:
    'create docx contract from dict with info -> output filename'
    # {
    #     "name": "НИУ ВШЭ",
    #     "signer": "директор Берсенев Илья Иванович",
    #     "document_name": "поставка усепешных программистов",
    #     "place": "г.Москва",
    #     "price": "1000"
    # }

    global n_contract
    input_filename = './docx_files/contract.docx' if input_filename is None else input_filename
    output_filename = f'.{dirname}/contract_{n_contract}.docx' if output_filename is None else output_filename
    print("############", input_filename, output_filename)
    print("############", dirname, os.path.exists(dirname))

    abs_p = '/Users/admin/Desktop/tender/tenderhack/back/chatapp/tenders'
    print(abs_p+dirname)
    if not os.path.exists(abs_p+dirname):
        os.mkdir(dirname)

    doc = Document(input_filename)
    raw_json['contract_n'] = n_contract
    raw_json['today'] = datetime.now().strftime('%d.%m.%Y')

    for k, v in raw_json.items():
        pattern = re.compile(r'\{' + re.escape(k) + r'=([^}]+)\}')

        for p in doc.paragraphs:
            p.text = pattern.sub('{' + str(k) + '=' + str(v) + '}', p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = pattern.sub('{' + str(k) + '=' + str(v) + '}', cell.text)

    n_contract += 1
    doc.save(output_filename)
    return filename


def file_to_docx(filename:str):
    "get clear file"
    doc = Document(filename)
    pattern = re.compile(r'\{([^}]+)=([^}]+)\}')

    for p in doc.paragraphs:
        p.text = re.sub(pattern, r'\2',  p.text)

    for table in doc.tables:
        for cell in table._cells:
            cell.text = re.sub(pattern, r'\2', cell.text)

    filename = f'{"".join(filename.split(".")[:-1])}_export.docx'
    doc.save(filename)
    return filename


def file_to_pdf(filename_docx:str, filename_pdf:str = None) -> str:
    'filename .docx -> convert file to .pdf'
    filename_pdf = '.' + ''.join(filename_docx.split('.')[:-1]) + '.pdf' if filename_pdf is None else filename_pdf
    convert(filename_docx,filename_pdf)
    return filename_pdf


def contract_change_value(raw_json:dict, filename:str):
    'tag value to value from json -> new_filename, report_name'
    # {
    # 'document_id': '1',
    # 'tag': 'place',
    # 'value': 'г. Москва',
    # 'comment' : 'НУЛЬ'
    # }
    doc = Document(filename)
    pattern_general = r'\{([^}]+)\}'
    matches = []
    old_values = []

    for p in doc.paragraphs:
        paragraph_matches = re.findall(pattern_general, p.text)
        matches += paragraph_matches
        old_values += [value for key, value in (pair.split('=') for pair in paragraph_matches) if
                       key == raw_json['tag']]
        p.text = re.sub(r'\{' + re.escape(raw_json["tag"]) + r'=([^}]+)\}',
                        '{' + raw_json["tag"] + '=' + raw_json["value"] + '}', p.text)

    # Iterate through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_matches = re.findall(pattern_general, cell.text)
                matches += cell_matches
                old_values += [value for key, value in (pair.split('=') for pair in cell_matches) if
                               key == raw_json['tag']]
                cell.text = re.sub(r'\{' + re.escape(raw_json["tag"]) + r'=([^}]+)\}',
                                   '{' + raw_json["tag"] + '=' + raw_json["value"] + '}', cell.text)

    # Save the modified document
    new_filename = f'{"".join(filename.split(".")[:-1])}_changed.docx'
    doc.save(new_filename)

    # Generate report
    all_values = {key: value for key, value in (pair.split('=') for pair in matches)}
    changes = [locations[raw_json['tag']], old_values[0], raw_json['value']]

    report_name = changes_report_fill(changes, all_values, old_values[0])

    # print("M", matches)
    # print("OLD ", *old_values)
    # print("NEW ", all_values[raw_json['tag']])
    return new_filename, report_name

def contract_change_punct(raw_json, filename):
    'tag value to value from json'
    # {
    #     'document_id': '1',
    #     'punct': '5.1.',
    #     'value': 'Мы вам все простим.',
    #     'comment': 'НУЛЬ'
    # }

    doc = Document(filename)
    for p in doc.paragraphs:
        if p.text.startswith(raw_json["punct"]):
            # print(p.text)
            p.text = raw_json["punct"] + ' ' + raw_json["value"]
            # print(p.text)

    # TODO : if punct does not exist - add next

    new_filename = f'{"".join(filename.split(".")[:-1])}_changed.docx'
    doc.save(new_filename)

    #gen report
    return new_filename


# def get_data_from_db(clause:ContractClause):
#     'input = tender'
#     tags = get_tags_from_docx()
#     data = dict()
#     supplier_tags = list(filter(lambda tag: 'supplier' in tag, tags))
#     contractor_tags = list(filter(lambda tag: 'contractor' in tag, tags))
#
#     data['clause_n'] = clause.tid
#     data['contract_n'] = clause.tender_id
#     data['today'] = datetime.now().strftime('%d.%m.%Y')
#
#     old = clause.before_clause.json()
#     new = clause.clause.json()
#
#
#     # IN PROGRESS
#
#     for tag in supplier_tags:
#         data[tag] = tag
#
#     for tag in contractor_tags:
#         data[tag] = tag



json1 = dict({
'document_id': '1',
'tag': 'place',
'value': 'г. Пермь',
'comment' : 'НУЛЬ'
})


json2 = dict({
'document_id': '1',
'punct': '4.1.',
'value': 'Мы вам все простим.',
'comment' : 'НУЛЬ'
})
# print(contract_change_punct("/Users/admin/Desktop/tender/tenderhack/back/docx_files/contract.docx", json2))


# FILL FROM FRONTEND
json3 = dict({
"name" : "ПЕРВЫЙ",
"signer" : "директор Берсенев Илья Иванович",
"document_name" : "поставка усепешных программистов",
"place" : "г.Владикавказ",
"price" : "1000"
})

# обработать signer + supplier
json3["supplier"] = json3["name"]
json3.pop("name")
json3["supplier_signer"] = json3["signer"]
json3.pop("signer")

# print(contract_fill(json3))


# FILL FROM FRONTEND
json3 = dict({
"name" : "ВТОРОЙ",
"signer" : "ВТОРОЙ ВТОРОЙ ВТОРОЙ ВТОРОЙ ",
"document_name" : "поставка усепешных программистов",
})

# обработать signer + supplier
json3["contractor"] = json3["name"]
json3.pop("name")
json3["contractor_signer"] = json3["signer"]
json3.pop("signer")

# print(contract_fill(json3, "/Users/admin/Desktop/tender/tenderhack/back/docx_files/contract.docx"))
# print(contract_change_value(json1, "/Users/admin/Desktop/tender/tenderhack/back/docx_files/contract_1.docx"))


# {
#     "numberField": "123",
#     "validityPeriod": {
#         "startDate": "2023-12-28T19:00:00.000Z",
#         "endDate": "2023-12-06T19:00:00.000Z"
#     },
#     "summ": "123",
#     "avans": "213",
#     "financeSource": "123",
#     "ikz": "123",
#     "place": "132",
#     "subject": "312"
# }

json4 = {
    "numberField": "123",
    "validityPeriod": {
        "startDate": "2023-12-28T19:00:00.000Z",
        "endDate": "2023-12-06T19:00:00.000Z"
    },
    "summ": "123",
    "avans": "213",
    "financeSource": "123",
    "ikz": "123",
    "place": "132",
    "subject": "312"
}
# print(contract_fill(json4))

