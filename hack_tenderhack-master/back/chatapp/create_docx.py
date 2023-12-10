import os
import re
from docx import Document
from docx2pdf import convert
from datetime import datetime
import constants
import json

# from .core import *
# from .db import engine, db_session, Base
# from .models.core import Tender, Document, ContractClause
# counter of report and contract
n_report = 1


# Протокол разногласий
def changes_report_fill(changes: list, raw_json: dict, dirname: str = '', old_value=None) -> str:
    'create changes report -> output filename'

    global n_report
    output_filename = f'.{dirname}/docx_files/changes_{n_report}.docx'

    if not os.path.exists(dirname):
        os.mkdir(dirname)

    doc = Document(input_filename)
    # changes = [locations[tag], '', raw_json[value]]
    changes_table = doc.tables[0]
    row = changes_table.add_row()
    row.cells[0].text = changes[0]
    row.cells[1].text = f"{old_value}"
    row.cells[2].text = changes[2]

    # raw_json['contract_n'] = n_contract
    # print(raw_json)
    pattern = r'\{"([^"]+)"="([^"]+)"\}'

    for p in doc.paragraphs:
        p.text = re.sub(pattern, r'\2', p.text)

    for table in doc.tables:
        for cell in table._cells:
            cell.text = re.sub(pattern, r'\2', cell.text)

    n_report += 1
    doc.save(output_filename)
    return filename


# Заполнение договора
def contract_fill(raw_json: dict, tender_id: str) -> str:
    'create docx contract from dict with info -> output filename'
    # {
    #     "name": "НИУ ВШЭ",
    #     "signer": "директор Берсенев Илья Иванович",
    #     "document_name": "поставка усепешных программистов",
    #     "place": "г.Москва",
    #     "price": "1000"
    # }

    # Путь к файлу-шаблону
    input_filename_path = constants.INPUT_FILENAME
    # Если нет папки для хранения создаём папку с названием id тендера
    tender_files_path = constants.TENDERS_DIR + tender_id

    if not os.path.exists(tender_files_path):
        os.mkdir(tender_files_path)
        n_contract = 1
    else:
        n_contract = len(os.listdir(tender_files_path))

    output_filename = f'{tender_files_path}/contract_{n_contract}.docx'

    # Создаём объект из шаблона
    doc = Document(input_filename_path)

    raw_json['contract_n'] = n_contract
    raw_json['today'] = datetime.now().strftime('%d.%m.%Y')

    for k, v in raw_json.items():
        pattern = re.compile(r'\{' + re.escape(k) + r'=([^}]+)\}')

        for p in doc.paragraphs:
            p.text = pattern.sub('{' + str(k) + '=' + str(v) + '}', p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = pattern.sub('{' + str(k) + '=' + str(v) + '}', cell.text)

    n_contract += 1
    doc.save(output_filename)
    print("FROM contract_fill", output_filename)
    return output_filename


def file_to_docx(filename: str):
    "get clear file"
    print("HELP!")
    print(filename)
    doc = Document(filename)
    pattern = re.compile(r'\{([^}]+)=([^}]+)\}')

    for p in doc.paragraphs:
        p.text = re.sub(pattern, r'\2', p.text)

    for table in doc.tables:
        for cell in table._cells:
            cell.text = re.sub(pattern, r'\2', cell.text)

    filename = f'{"".join(filename.split(".")[:-1])}_export.docx'
    doc.save(filename)
    return filename


def file_to_pdf(filename_docx: str, filename_pdf: str = None) -> str:
    'filename .docx -> convert file to .pdf'
    filename_pdf = '.' + ''.join(filename_docx.split('.')[:-1]) + '.pdf' if filename_pdf is None else filename_pdf
    convert(filename_docx, filename_pdf)
    return filename_pdf


# Меняет пункт в договоре
# Создаёт протокол разногласий и туда заносит изменённый пункт
def contract_change_value(raw_json: dict, filename: str):
    'tag value to value from json -> new_filename, report_name'
    # {
    # 'document_id': '1',
    # 'tag': 'place',
    # 'value': 'г. Москва',
    # 'comment' : 'НУЛЬ'
    # }
    doc = Document(filename)
    pattern_general = r'\{([^}]+)\}'
    matches = []
    old_values = []

    for p in doc.paragraphs:
        paragraph_matches = re.findall(pattern_general, p.text)
        matches += paragraph_matches
        old_values += [value for key, value in (pair.split('=') for pair in paragraph_matches) if
                       key == raw_json['tag']]
        p.text = re.sub(r'\{' + re.escape(raw_json["tag"]) + r'=([^}]+)\}',
                        '{' + raw_json["tag"] + '=' + raw_json["value"] + '}', p.text)

    # Iterate through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_matches = re.findall(pattern_general, cell.text)
                matches += cell_matches
                old_values += [value for key, value in (pair.split('=') for pair in cell_matches) if
                               key == raw_json['tag']]
                cell.text = re.sub(r'\{' + re.escape(raw_json["tag"]) + r'=([^}]+)\}',
                                   '{' + raw_json["tag"] + '=' + raw_json["value"] + '}', cell.text)

    # Save the modified document
    new_filename = f'{"".join(filename.split(".")[:-1])}_changed.docx'
    doc.save(new_filename)

    # Generate report
    all_values = {key: value for key, value in (pair.split('=') for pair in matches)}
    changes = [locations[raw_json['tag']], old_values[0], raw_json['value']]

    report_name = changes_report_fill(changes, all_values, old_values[0])

    # print("M", matches)
    # print("OLD ", *old_values)
    # print("NEW ", all_values[raw_json['tag']])
    return new_filename, report_name


def contract_change_punct(raw_json, filename):
    'tag value to value from json'
    # {
    #     'document_id': '1',
    #     'punct': '5.1.',
    #     'value': 'Мы вам все простим.',
    #     'comment': 'НУЛЬ'
    # }

    doc = Document(filename)
    for p in doc.paragraphs:
        if p.text.startswith(raw_json["punct"]):
            # print(p.text)
            p.text = raw_json["punct"] + ' ' + raw_json["value"]
            # print(p.text)

    # TODO : if punct does not exist - add next

    new_filename = f'{"".join(filename.split(".")[:-1])}_changed.docx'
    doc.save(new_filename)

    # gen report
    return new_filename


json1 = dict({
    'document_id': '1',
    'tag': 'place',
    'value': 'г. Пермь',
    'comment': 'НУЛЬ'
})

json2 = dict({
    'document_id': '1',
    'punct': '4.1.',
    'value': 'Мы вам все простим.',
    'comment': 'НУЛЬ'
})
# print(contract_change_punct("/Users/admin/Desktop/tender/tenderhack/back/docx_files/contract.docx", json2))


# FILL FROM FRONTEND
json3 = dict({
    "name": "ПЕРВЫЙ",
    "signer": "директор Берсенев Илья Иванович",
    "document_name": "поставка усепешных программистов",
    "place": "г.Владикавказ",
    "price": "1000"
})

# обработать signer + supplier
json3["supplier"] = json3["name"]
json3.pop("name")
json3["supplier_signer"] = json3["signer"]
json3.pop("signer")

# print(contract_fill(json3))


# FILL FROM FRONTEND
json3 = dict({
    "name": "ВТОРОЙ",
    "signer": "ВТОРОЙ ВТОРОЙ ВТОРОЙ ВТОРОЙ ",
    "document_name": "поставка усепешных программистов",
})

# обработать signer + supplier
json3["contractor"] = json3["name"]
json3.pop("name")
json3["contractor_signer"] = json3["signer"]
json3.pop("signer")

# print(contract_fill(json3, "/Users/admin/Desktop/tender/tenderhack/back/docx_files/contract.docx"))
# print(contract_change_value(json1, "/Users/admin/Desktop/tender/tenderhack/back/docx_files/contract_1.docx"))


# {
#     "numberField": "123",
#     "validityPeriod": {
#         "startDate": "2023-12-28T19:00:00.000Z",
#         "endDate": "2023-12-06T19:00:00.000Z"
#     },
#     "summ": "123",
#     "avans": "213",
#     "financeSource": "123",
#     "ikz": "123",
#     "place": "132",
#     "subject": "312"
# }

json4 = {
    "numberField": "123",
    "validityPeriod": {
        "startDate": "2023-12-28T19:00:00.000Z",
        "endDate": "2023-12-06T19:00:00.000Z"
    },
    "summ": "123",
    "avans": "213",
    "financeSource": "123",
    "ikz": "123",
    "place": "132",
    "subject": "312"
}
# print(contract_fill(json4))
